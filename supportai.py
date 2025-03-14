import os
import sys
import signal
import argparse
import ollama
import PyPDF2
import docx
import pandas as pd
from pathlib import Path

# Graceful exit on Ctrl+C
def signal_handler(sig, frame):
    print("\nAborted by the user.")
    sys.exit(0)

signal.signal(signal.SIGINT, signal_handler)

# Read folder path from folder.txt
def get_folder_path():
    try:
        with open("folder.txt", "r") as f:
            return f.read().strip()
    except FileNotFoundError:
        print("Error: folder.txt not found.")
        sys.exit(1)
    except Exception as e:
        print(f"Error reading folder.txt: {e}")
        sys.exit(1)

# Read text from different file formats efficiently
def read_file_content(file_path):
    ext = file_path.suffix.lower()
    try:
        if ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        elif ext == ".docx":
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + " | ".join(cell.text for cell in row.cells)
            return text
        elif ext in [".xlsx", ".csv"]:
            df = pd.read_excel(file_path) if ext == ".xlsx" else pd.read_csv(file_path)
            return df.to_string()
        elif ext == ".pdf":
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
            return text
        else:
            return None
    except Exception as e:
        print(f"Error reading {file_path.name}: {e}")
        return None

# Load all file contents efficiently
def load_files(folder):
    folder_path = Path(folder)
    if not folder_path.exists() or not folder_path.is_dir():
        print("Error: Specified folder does not exist.")
        sys.exit(1)
    
    content = []
    file_stats = []
    for file in folder_path.iterdir():
        file_text = read_file_content(file)
        if file_text:
            char_count = len(file_text)
            word_count = len(file_text.split())
            file_stats.append((file.name, char_count, word_count))
            content.append(f"\n\n--- {file.name} ---\n{file_text}")
    
    if not content:
        print("⚠️ No readable content found in files.")
        return "No readable files found."
    
    full_content = "\n".join(content)
    print(f"✅ Loaded {len(full_content)} characters from {len(file_stats)} files")  # Debugging output
    print("\nFile Statistics:")
    for file_name, char_count, word_count in file_stats:
        print(f"{file_name}: {char_count} characters, {word_count} words")
    
    return full_content

# Main AI interaction loop
def main(ollama_model):
    folder = get_folder_path()
    file_data = load_files(folder)
    
    print("\nAI Assistant is ready. Type your query or 'exit' to quit.")
    
    while True:
        try:
            user_input = input("AI Assistant: ")
            if user_input.lower() in ["exit", "bye"]:
                print("Goodbye!")
                break
            
            # Send file content in chunks to avoid overloading the model
            chunk_size = 5000  # Adjust based on model capabilities
            chunks = [file_data[i:i+chunk_size] for i in range(0, len(file_data), chunk_size)]
            responses = []
            
            for chunk in chunks:
                response = ollama.chat(model=ollama_model, messages=[
                    {"role": "system", "content": "Analyze the given documents and respond based on user queries."},
                    {"role": "user", "content": chunk},
                    {"role": "user", "content": user_input}
                ])
                responses.append(response["message"]["content"])
            
            print("\nResponse:\n", "\n".join(responses), "\n")
        
        except Exception as e:
            print(f"Error: {e}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("-l", required=True, help="Specify Ollama AI model")
    args = parser.parse_args()
    main(args.l)
